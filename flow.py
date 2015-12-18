import ConfigParser
config = ConfigParser.ConfigParser()
config.read("flow.config")

import datetime
from docx import Document
from docx.shared import Inches
document = Document()

document.add_heading('Case ' + config.get('Flow','Case'), level=1)
document.add_paragraph('\nClinical History:\n'+ config.get('Flow','ClinicalHistory'))

document.add_paragraph('Hematological Value:\nWBC='+ config.get('Flow','HemeWBC')+"k/uL, Hemoglobin="+ config.get('Flow','HemeHGB')+"g/dL, Hematocrit="+ config.get('Flow','HemeHCT')+"%, Platelet="+ config.get('Flow','HemePLT'))

document.add_paragraph('Cytospin: \n'+ config.get('Flow','Cytospin'))

interp = config.get('Flow','Interpretation')
s=[]
items = config.items('IHC');
items.sort()
for (i,j) in items[0:len(items)-1]:
    s.append(i.upper())
t = ", ".join(s)
(i,j) = items[len(items)-1]
t += " and " + i.upper()

document.add_paragraph('Interpretation: \n' + interp.replace("_IHC_",t))

document.add_paragraph('Diagnosis: \n' +config.get('Flow','Diagnosis'))

document.add_paragraph('Pathologist: '+config.get('Flow','Pathologist'))

today = datetime.date.today()
formatdate = str(today.month)+"/"+str(today.day)+"/"+str(today.year)
document.add_paragraph('Date: ' + formatdate)

#document.add_page_break()
document.add_paragraph('-------------------------------------------------')

document.add_paragraph('Strategy: '+config.get('Flow','Strategy'))

mylist = config.get('Form','List')
mylist = mylist.split(",")
s = ""
for i in mylist:
    try:
        t = config.get('IHC',i)
        ti = float(t)
        s += i + ": %0.0f\n" % ti
    except ConfigParser.NoOptionError:
        s += i + ": \n"
document.add_paragraph(s)

document.save('flow.docx')





